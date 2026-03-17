"""
Module 1: Document text extraction.
Supports PDF (digital + scanned via OCR) and DOCX files.
"""

import os
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber


def extract_text(file_path: str) -> str:
    """Extract text from PDF or DOCX. Falls back to OCR for scanned PDFs."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Bestand niet gevonden: {file_path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Onbekend bestandstype: {ext} (gebruik .pdf of .docx)")


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF. Uses pdfplumber first, falls back to PyMuPDF, then OCR."""
    # Strategy 1: pdfplumber (best for structured/table PDFs)
    text = _extract_with_pdfplumber(file_path)
    if _has_meaningful_text(text):
        return text

    # Strategy 2: PyMuPDF (good for general PDFs)
    text = _extract_with_pymupdf(file_path)
    if _has_meaningful_text(text):
        return text

    # Strategy 3: OCR via Tesseract (for scanned documents)
    text = _extract_with_ocr(file_path)
    if _has_meaningful_text(text):
        return text

    raise ValueError("Kon geen tekst extraheren uit het PDF-bestand. Is het leeg of beschadigd?")


def _extract_with_pdfplumber(file_path: str) -> str:
    """Extract using pdfplumber — good for forms and tables."""
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception:
        return ""


def _extract_with_pymupdf(file_path: str) -> str:
    """Extract using PyMuPDF — good general-purpose extractor."""
    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)
    except Exception:
        return ""


def _extract_with_ocr(file_path: str) -> str:
    """Extract using Tesseract OCR — for scanned/image-based PDFs."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page as image at 300 DPI
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            # OCR with Dutch language
            page_text = pytesseract.image_to_string(img, lang="nld")
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()
        return "\n\n".join(text_parts)
    except Exception:
        return ""


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is vereist voor DOCX-bestanden: pip install python-docx")

    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def _has_meaningful_text(text: str, min_chars: int = 50) -> bool:
    """Check if extracted text has enough meaningful content."""
    if not text:
        return False
    # Strip whitespace and check length
    cleaned = text.strip()
    if len(cleaned) < min_chars:
        return False
    # Check it's not just garbage characters
    alpha_ratio = sum(c.isalpha() for c in cleaned) / max(len(cleaned), 1)
    return alpha_ratio > 0.3
